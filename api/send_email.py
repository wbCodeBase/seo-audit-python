"""
/api/send_email  POST {job_id, email}
Sends:
  1. Summary HTML email body
  2. Full SEO report as a .docx Word attachment

Required Vercel env vars:
  SMTP_HOST  SMTP_PORT  SMTP_USER  SMTP_PASS  SMTP_FROM
"""
import json, os, re, smtplib, ssl, io
from email.mime.multipart import MIMEMultipart
from email.mime.text      import MIMEText
from email.mime.base      import MIMEBase
from email                import encoders
from http.server          import BaseHTTPRequestHandler


# ── Redis ─────────────────────────────────────────────────────────────────────

def get_redis():
    url   = os.environ.get("UPSTASH_REDIS_REST_URL")
    token = os.environ.get("UPSTASH_REDIS_REST_TOKEN")
    if not url or not token:
        return None
    try:
        from upstash_redis import Redis
        return Redis(url=url, token=token)
    except Exception:
        return None

def store_get(job_id):
    r = get_redis()
    if not r:
        return None
    try:
        v = r.get(f"seo:{job_id}")
        return json.loads(v) if v else None
    except Exception:
        return None

def valid_email(addr):
    return bool(re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]{2,}$', addr.strip()))


# ── Summary email body (HTML) ─────────────────────────────────────────────────

def build_summary_html(D, name, report_email):
    domain = D.get("domain", "your website")
    ov     = D.get("overall", {})
    grade  = ov.get("grade", "—")
    summary= ov.get("summary", "")
    cats   = D.get("cats", [])
    recs   = D.get("recommendations", [])
    op     = D.get("op", {})
    gc = ("#1e8449" if grade in ["A+","A","A-","B+"] else
          "#b7770d" if grade in ["B","B-","C+","C"] else "#c0392b")
    cat_rows = "".join(
        f'<tr><td style="padding:8px 12px;font-size:13px;color:#1c2b3a">{c.get("lbl","")}</td>'
        f'<td style="padding:8px 12px;font-size:13px;font-weight:700;color:{gc}">{c.get("grade","—")}</td></tr>'
        for c in cats)
    rec_rows = "".join(
        f'<tr><td style="padding:6px 0;vertical-align:top;width:24px">'
        f'<span style="display:inline-block;width:20px;height:20px;border-radius:50%;'
        f'background:#b7770d;color:#fff;font-size:10px;font-weight:700;text-align:center;line-height:20px">'
        f'{r.get("priority",i+1)}</span></td>'
        f'<td style="padding:6px 0 6px 10px">'
        f'<div style="font-size:13px;font-weight:600;color:#1c2b3a;margin-bottom:3px">{r.get("title","")}</div>'
        f'<div style="font-size:12px;color:#5d6d7e;line-height:1.6">{r.get("detail","")}</div></td></tr>'
        for i, r in enumerate(recs[:6]))
    title_t = (op.get("title") or {}).get("t","—")
    meta_t  = (op.get("meta")  or {}).get("t","—")
    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"/></head>
<body style="margin:0;padding:0;background:#eef0f4;font-family:'Segoe UI',Arial,sans-serif">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#eef0f4;padding:32px 16px">
<tr><td align="center"><table width="600" cellpadding="0" cellspacing="0" style="max-width:600px;width:100%">
<tr><td style="background:linear-gradient(135deg,#c0392b,#922b21);border-radius:16px 16px 0 0;padding:30px 34px;color:#fff">
  <div style="font-size:22px;font-weight:700;margin-bottom:6px">SEO Report for {domain}</div>
  <div style="font-size:13px;opacity:.8">Prepared for {name}</div>
</td></tr>
<tr><td style="background:#fff;padding:28px 34px;border-bottom:1px solid #eee">
  <table width="100%" cellpadding="0" cellspacing="0"><tr>
    <td style="vertical-align:middle">
      <div style="font-size:11px;font-weight:600;text-transform:uppercase;color:#95a5a6;margin-bottom:6px">Overall Grade</div>
      <div style="font-size:52px;font-weight:700;color:{gc};line-height:1">{grade}</div>
    </td>
    <td style="vertical-align:middle;padding-left:28px">
      <div style="font-size:13.5px;color:#1c2b3a;line-height:1.65">{summary}</div>
    </td>
  </tr></table>
</td></tr>
<tr><td style="background:#fff;padding:0 34px 24px">
  <div style="font-size:11px;font-weight:600;text-transform:uppercase;color:#95a5a6;margin-bottom:12px">Category Scores</div>
  <table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #eee;border-radius:10px;overflow:hidden">
    <thead><tr style="background:#f8f9fb">
      <th style="padding:8px 12px;text-align:left;font-size:10px;text-transform:uppercase;color:#95a5a6;font-weight:600">Category</th>
      <th style="padding:8px 12px;text-align:left;font-size:10px;text-transform:uppercase;color:#95a5a6;font-weight:600">Grade</th>
    </tr></thead><tbody>{cat_rows}</tbody>
  </table>
</td></tr>
<tr><td style="background:#fff;padding:0 34px 24px;border-bottom:1px solid #eee">
  <div style="font-size:11px;font-weight:600;text-transform:uppercase;color:#95a5a6;margin-bottom:12px">Key Findings</div>
  <table width="100%" cellpadding="0" cellspacing="0">
    <tr><td style="padding:8px 0;border-bottom:1px solid #f0f2f5">
      <div style="font-size:11px;color:#95a5a6;margin-bottom:3px">TITLE TAG</div>
      <div style="font-size:13px;color:#1c2b3a">{(title_t[:80]+"...") if len(title_t)>80 else title_t}</div>
    </td></tr>
    <tr><td style="padding:8px 0">
      <div style="font-size:11px;color:#95a5a6;margin-bottom:3px">META DESCRIPTION</div>
      <div style="font-size:13px;color:#1c2b3a">{(meta_t[:120]+"...") if len(meta_t)>120 else meta_t}</div>
    </td></tr>
  </table>
</td></tr>
<tr><td style="background:#fff;padding:24px 34px;border-bottom:1px solid #eee">
  <div style="font-size:11px;font-weight:600;text-transform:uppercase;color:#95a5a6;margin-bottom:14px">🎯 Priority Recommendations</div>
  <table width="100%" cellpadding="0" cellspacing="0"><tbody>{rec_rows}</tbody></table>
</td></tr>
<tr><td style="background:#fff;padding:20px 34px;text-align:center;border-bottom:1px solid #eee">
  <div style="font-size:13px;color:#5d6d7e">
    📎 <strong>Full report attached</strong> as a Word document (.docx)<br/>
    Open in Microsoft Word or Google Docs to view all sections.
  </div>
</td></tr>
<tr><td style="background:#f8f9fb;border-radius:0 0 16px 16px;padding:18px 34px">
  <div style="font-size:12px;color:#95a5a6">
    SEO Audit for <strong style="color:#c0392b">{domain}</strong> · Powered by Claude AI<br/>
    Sent to {report_email}
  </div>
</td></tr>
</table></td></tr></table>
</body></html>"""


def build_summary_text(D, name):
    domain  = D.get("domain","your website")
    ov      = D.get("overall",{})
    cats    = D.get("cats",[])
    recs    = D.get("recommendations",[])
    cat_str = "\n".join(f"  {c.get('lbl')}: {c.get('grade')}" for c in cats)
    rec_str = "\n".join(f"  {i+1}. {r.get('title')}\n     {r.get('detail')}" for i,r in enumerate(recs[:6]))
    return f"""SEO Audit Report for {domain}
Prepared for {name}

Overall Grade: {ov.get('grade','—')}
{ov.get('summary','')}

Category Scores:
{cat_str}

Priority Recommendations:
{rec_str}

---
Full report attached as a Word document (.docx).
Open in Microsoft Word or Google Docs.
Powered by Claude AI SEO Audit Tool
"""


# ── Build DOCX report ─────────────────────────────────────────────────────────

def build_docx(D, name, email, date):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text  import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns    import qn
    from docx.oxml       import OxmlElement

    GRADES = ["A+","A","A-","B+","B","B-","C+","C","C-","D+","D","D-","F"]
    def gp(g): i=GRADES.index(g) if g in GRADES else 6; return 1-(i/len(GRADES))
    def gc_rgb(g):
        p=gp(g)
        if p>.75: return RGBColor(0x1E,0x84,0x49)
        if p>.5:  return RGBColor(0xB7,0x77,0x0D)
        if p>.3:  return RGBColor(0xE6,0x7E,0x22)
        return RGBColor(0xC0,0x39,0x2B)

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def set_table_borders(table, color='DDDDDD'):
        for row in table.rows:
            for cell in row.cells:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcB  = OxmlElement('w:tcBorders')
                for edge in ('top','bottom','left','right','insideH','insideV'):
                    tag = OxmlElement(f'w:{edge}')
                    tag.set(qn('w:val'),   'single')
                    tag.set(qn('w:sz'),    '4')
                    tag.set(qn('w:color'), color)
                    tcB.append(tag)
                tcPr.append(tcB)

    def add_run(para, text, bold=False, italic=False, size=11,
                color=None, font='Arial'):
        r = para.add_run(text)
        r.bold       = bold
        r.italic     = italic
        r.font.size  = Pt(size)
        r.font.name  = font
        if color:
            r.font.color.rgb = color
        return r

    def section_heading(doc, title, bg_hex, text_hex='FFFFFF'):
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        cell = t.rows[0].cells[0]
        set_cell_bg(cell, bg_hex)
        cell.width = Inches(6.5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.left_indent  = Pt(6)
        rgb = RGBColor(int(text_hex[0:2],16),int(text_hex[2:4],16),int(text_hex[4:6],16))
        add_run(p, title, bold=True, size=13, color=rgb)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def check_row(doc, status, label, detail=''):
        sym = {'p':'✓ ','f':'✗ ','w':'! ','i':'ℹ '}[status]
        col = {'p':RGBColor(0x1E,0x84,0x49),'f':RGBColor(0xC0,0x39,0x2B),
               'w':RGBColor(0xB7,0x77,0x0D),'i':RGBColor(0x1A,0x52,0x76)}[status]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        add_run(p, sym, bold=True, size=11, color=col)
        add_run(p, label, bold=False, size=11)
        if detail:
            pd = doc.add_paragraph()
            pd.paragraph_format.left_indent  = Cm(1.2)
            pd.paragraph_format.space_before = Pt(0)
            pd.paragraph_format.space_after  = Pt(4)
            add_run(pd, detail, size=10, color=RGBColor(0x5D,0x6D,0x7E))

    # ── Data ──────────────────────────────────────────────────────────────────
    domain = D.get("domain","")
    ov     = D.get("overall",{})
    cats   = D.get("cats",[])
    op     = D.get("op",{})
    geo    = D.get("geo",{})
    us     = D.get("us",{})
    pf     = D.get("pf",{})
    tech   = D.get("tech",{})
    loc    = D.get("local",{})
    recs   = D.get("recommendations",[])
    cwv    = us.get("cwv",{})
    mob    = us.get("mob",{})
    desk   = us.get("desk",{})
    spd    = pf.get("speed",{})
    sz     = pf.get("size",{})
    gbp    = loc.get("gbp",{})
    rev    = loc.get("reviews",{})

    ov_grade = ov.get("grade","B")
    ov_color = gc_rgb(ov_grade)

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2)
    sec.right_margin  = Cm(2)

    # ── HEADER ────────────────────────────────────────────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=1)
    hdr_tbl.style = 'Table Grid'
    hcell = hdr_tbl.rows[0].cells[0]
    set_cell_bg(hcell, 'C0392B')
    hp = hcell.paragraphs[0]
    hp.paragraph_format.space_before = Pt(14)
    hp.paragraph_format.space_after  = Pt(4)
    hp.paragraph_format.left_indent  = Pt(8)
    add_run(hp, f'SEO Audit Report — {domain}', bold=True, size=17,
            color=RGBColor(0xFF,0xFF,0xFF))
    hp2 = hcell.add_paragraph()
    hp2.paragraph_format.space_before = Pt(0)
    hp2.paragraph_format.space_after  = Pt(12)
    hp2.paragraph_format.left_indent  = Pt(8)
    add_run(hp2, f'Prepared for {name}  ·  {date}', size=10,
            color=RGBColor(0xFF,0xCC,0xCC))
    doc.add_paragraph()

    # ── OVERALL GRADE ─────────────────────────────────────────────────────────
    ov_tbl = doc.add_table(rows=1, cols=2)
    ov_tbl.style = 'Table Grid'
    set_table_borders(ov_tbl, 'EEEEEE')

    # Left cell: grade
    lc = ov_tbl.rows[0].cells[0]
    lc.width = Inches(1.4)
    set_cell_bg(lc, 'F8F9FB')
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(10)
    lp.paragraph_format.space_after  = Pt(4)
    add_run(lp, 'OVERALL GRADE', size=8,
            color=RGBColor(0x95,0xA5,0xA6))
    lp2 = lc.add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(10)
    add_run(lp2, ov_grade, bold=True, size=36, color=ov_color)

    # Right cell: summary
    rc = ov_tbl.rows[0].cells[1]
    rp = rc.paragraphs[0]
    rp.paragraph_format.space_before = Pt(10)
    rp.paragraph_format.space_after  = Pt(4)
    rp.paragraph_format.left_indent  = Pt(8)
    add_run(rp, 'Summary', bold=True, size=10,
            color=RGBColor(0x5D,0x6D,0x7E))
    rp2 = rc.add_paragraph()
    rp2.paragraph_format.space_before = Pt(2)
    rp2.paragraph_format.space_after  = Pt(10)
    rp2.paragraph_format.left_indent  = Pt(8)
    add_run(rp2, ov.get("summary",""), size=11)

    doc.add_paragraph()

    # ── CATEGORY SCORES ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run(p, 'Category Breakdown', bold=True, size=12)
    p.paragraph_format.space_after = Pt(4)

    cat_tbl = doc.add_table(rows=1, cols=2)
    cat_tbl.style = 'Table Grid'
    set_table_borders(cat_tbl, 'DDDDDD')
    # Header
    for i, hd in enumerate(['Category','Grade']):
        cell = cat_tbl.rows[0].cells[i]
        set_cell_bg(cell, 'F8F9FB')
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after  = Pt(4)
        hp.paragraph_format.left_indent  = Pt(6)
        add_run(hp, hd, bold=True, size=9,
                color=RGBColor(0x95,0xA5,0xA6))
    # Data rows
    for cat in cats:
        row = cat_tbl.add_row()
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.left_indent  = Pt(6)
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        add_run(lp, cat.get('lbl',''), size=11)
        rp = row.cells[1].paragraphs[0]
        rp.paragraph_format.left_indent  = Pt(6)
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after  = Pt(4)
        add_run(rp, cat.get('grade','—'), bold=True, size=13,
                color=gc_rgb(cat.get('grade','B')))

    doc.add_paragraph()

    # ── ON-PAGE SEO ───────────────────────────────────────────────────────────
    section_heading(doc, '🔍  On-Page SEO Results', 'C0392B')

    title_t   = (op.get('title') or {}).get('t','')
    title_len = (op.get('title') or {}).get('len',0)
    title_ok  = 45 <= title_len <= 65
    meta_t    = (op.get('meta') or {}).get('t','')
    meta_len  = (op.get('meta') or {}).get('len',0)
    meta_ok   = 120 <= meta_len <= 165

    # SERP preview table
    serp = doc.add_table(rows=3, cols=1)
    serp.style = 'Table Grid'
    set_table_borders(serp, 'CCCCCC')
    serp_data = [
        ('URL',   op.get('serpUrl',f'https://{domain}'), RGBColor(0x18,0x80,0x38)),
        ('Title', op.get('serpTitle',title_t[:57]),       RGBColor(0x15,0x58,0xD6)),
        ('Desc',  op.get('serpDesc',meta_t[:155]),        RGBColor(0x3C,0x40,0x43)),
    ]
    for i,(lbl,val,col) in enumerate(serp_data):
        cell = serp.rows[i].cells[0]
        set_cell_bg(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent  = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run(p, f'{lbl}: ', bold=True, size=9,
                color=RGBColor(0x95,0xA5,0xA6))
        add_run(p, val, size=10, color=col)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    check_row(doc, 'p' if title_ok else 'w',
              f'Title Tag — {title_len} chars — {"optimal" if title_ok else "adjust to 50–60 chars"}',
              title_t)
    check_row(doc, 'p' if meta_ok else 'w',
              f'Meta Description — {meta_len} chars — {"optimal" if meta_ok else "adjust to 120–160 chars"}',
              meta_t[:150])

    # H1
    h1s = op.get('h1',[])
    h1_ok = len(h1s)==1
    check_row(doc, 'p' if h1_ok else 'w',
              f'H1 Tag — {"one H1 found" if h1_ok else f"{len(h1s)} H1 tags"}',
              h1s[0].get('v','') if h1s else 'Not found')

    # Keywords table
    kws = op.get('kws',[])
    if kws:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        add_run(p, 'Keyword Consistency', bold=True, size=10)
        kw_tbl = doc.add_table(rows=1, cols=5)
        kw_tbl.style = 'Table Grid'
        set_table_borders(kw_tbl, 'DDDDDD')
        for i, hd in enumerate(['Keyword','In Title','In Meta','In Headings','Frequency']):
            cell = kw_tbl.rows[0].cells[i]
            set_cell_bg(cell, 'F8F9FB')
            hp = cell.paragraphs[0]
            hp.paragraph_format.left_indent = Pt(4)
            hp.paragraph_format.space_before = Pt(3)
            hp.paragraph_format.space_after  = Pt(3)
            add_run(hp, hd, bold=True, size=8, color=RGBColor(0x95,0xA5,0xA6))
        for kw in kws:
            row = kw_tbl.add_row()
            vals = [kw.get('p',''), '✓' if kw.get('ti') else '—',
                    '✓' if kw.get('me') else '—',
                    '✓' if kw.get('hd') else '—',
                    str(kw.get('f','—'))]
            for i, val in enumerate(vals):
                p = row.cells[i].paragraphs[0]
                p.paragraph_format.left_indent  = Pt(4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                c = RGBColor(0x1E,0x84,0x49) if val=='✓' else None
                add_run(p, val, size=10, bold=(i==0), color=c)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    wc = op.get('wc',0)
    check_row(doc, 'p' if op.get('wcOk') else 'w',
              f'Word Count — {wc:,} words — {"good" if op.get("wcOk") else "below 500 minimum"}')
    check_row(doc, 'p' if op.get('imgAlt') else 'w',
              f'Image Alt Attributes — {op.get("imgAltDesc","")}')
    check_row(doc, 'p' if op.get('canonOk') else 'w',
              f'Canonical Tag — {"configured" if op.get("canonOk") else "missing"}',
              op.get('canon','Not detected'))
    check_row(doc, 'p' if op.get('noindexOk') else 'f',
              f'Noindex — {"page is indexable" if op.get("noindexOk") else "BLOCKING this page!"}')
    check_row(doc, 'p' if op.get('httpsRedir') else 'f',
              f'HTTPS — {"secure" if op.get("httpsRedir") else "not secure"}')
    check_row(doc, 'p' if op.get('robotsOk') else 'f',
              f'Robots.txt — {"found" if op.get("robotsOk") else "missing"}',
              op.get('robots',''))
    check_row(doc, 'p' if op.get('sitemapOk') else 'f',
              f'XML Sitemap — {"found" if op.get("sitemapOk") else "missing"}',
              op.get('sitemap',''))
    check_row(doc, 'p' if op.get('analytics') else 'w',
              f'Analytics — {"detected" if op.get("analytics") else "not detected"}',
              ', '.join(op.get('analyticsTools',[]) or ['None']))
    check_row(doc, 'p' if op.get('schema') else 'w',
              f'Schema.org — {"detected" if op.get("schema") else "not found"}',
              'Types: '+', '.join(op.get('schemaTypes',[])) if op.get('schema') else 'Add JSON-LD schema markup.')
    doc.add_paragraph()

    # ── GEO / AI ──────────────────────────────────────────────────────────────
    section_heading(doc, '🌐  GEO / AI Visibility', '1E8449')
    check_row(doc, 'p' if geo.get('renderOk') else 'w',
              f'LLM Readability — Render %: {geo.get("renderPct","N/A")}',
              geo.get('renderDesc',''))
    check_row(doc, 'p' if geo.get('llmsTxt') else 'w',
              f'llms.txt — {"found" if geo.get("llmsTxt") else "not found"}',
              geo.get('llmsTxtUrl',''))

    traffic = geo.get('traffic',{})
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    add_run(p, f'Traffic: Organic {traffic.get("org",0):,}  ·  Paid {traffic.get("paid",0):,}  ·  AI Overviews {traffic.get("ai",0):,}',
            size=10, color=RGBColor(0x5D,0x6D,0x7E))
    doc.add_paragraph()

    # ── USABILITY ─────────────────────────────────────────────────────────────
    section_heading(doc, '📱  Usability', 'C0392B')
    check_row(doc, 'p' if cwv.get('pass') else 'f',
              f'Core Web Vitals — {"passed" if cwv.get("pass") else "failed"}',
              f'LCP: {cwv.get("lcp","N/A")}  ·  INP: {cwv.get("inp","N/A")}  ·  CLS: {cwv.get("cls","N/A")}')
    ms = mob.get('score',0)
    check_row(doc, 'p' if ms>=70 else 'w' if ms>=50 else 'f',
              f'Mobile PageSpeed — {ms}/100',
              f'FCP: {mob.get("fcp","N/A")}  ·  LCP: {mob.get("lcp","N/A")}  ·  TTI: {mob.get("tti","N/A")}')
    ds = desk.get('score',0)
    check_row(doc, 'p' if ds>=70 else 'w' if ds>=50 else 'f',
              f'Desktop PageSpeed — {ds}/100',
              f'FCP: {desk.get("fcp","N/A")}  ·  LCP: {desk.get("lcp","N/A")}  ·  TTI: {desk.get("tti","N/A")}')
    check_row(doc, 'p' if us.get('viewport') else 'w',
              f'Mobile Viewport — {"specified" if us.get("viewport") else "not specified"}')
    check_row(doc, 'w' if us.get('iframes') else 'p',
              f'iFrames — {"detected" if us.get("iframes") else "none detected"}',
              us.get('iframesDesc',''))
    check_row(doc, 'p' if us.get('favicon') else 'w',
              f'Favicon — {"found" if us.get("favicon") else "not found"}')
    doc.add_paragraph()

    # ── PERFORMANCE ───────────────────────────────────────────────────────────
    section_heading(doc, '⚡  Performance', '1A5276')
    check_row(doc, 'p' if spd.get('ok') else 'w',
              'Website Load Speed',
              f'Server: {spd.get("srv","N/A")}  ·  Content: {spd.get("cnt","N/A")}  ·  Scripts: {spd.get("scr","N/A")}')
    check_row(doc, 'p' if sz.get('ok') else 'w',
              f'Download Size — {sz.get("tot","N/A")}',
              f'HTML: {sz.get("html","N/A")}  ·  CSS: {sz.get("css","N/A")}  ·  JS: {sz.get("js","N/A")}  ·  Images: {sz.get("img","N/A")}')
    check_row(doc, 'p' if pf.get('http2') else 'w',
              f'HTTP/2 — {"in use" if pf.get("http2") else "not in use"}')
    check_row(doc, 'f' if pf.get('jsErrors') else 'p',
              f'JavaScript Errors — {"detected" if pf.get("jsErrors") else "none"}',
              pf.get('jsErrDesc',''))
    doc.add_paragraph()

    # ── SOCIAL ────────────────────────────────────────────────────────────────
    section_heading(doc, '📣  Social Results', '6C3483')
    social = D.get('social',[])
    if social:
        soc_tbl = doc.add_table(rows=1, cols=3)
        soc_tbl.style = 'Table Grid'
        set_table_borders(soc_tbl, 'EEEEEE')
        for i,hd in enumerate(['Platform','Status','URL']):
            cell = soc_tbl.rows[0].cells[i]
            set_cell_bg(cell,'F8F9FB')
            hp = cell.paragraphs[0]
            hp.paragraph_format.left_indent = Pt(4)
            hp.paragraph_format.space_before = Pt(3)
            hp.paragraph_format.space_after  = Pt(3)
            add_run(hp, hd, bold=True, size=8, color=RGBColor(0x95,0xA5,0xA6))
        for s in social:
            row = soc_tbl.add_row()
            vals = [s.get('name',''), '✓ Linked' if s.get('linked') else '— Not linked', s.get('url','')]
            colors = [None, RGBColor(0x1E,0x84,0x49) if s.get('linked') else RGBColor(0xCC,0xCC,0xCC), RGBColor(0x5D,0x6D,0x7E)]
            for i,(val,col) in enumerate(zip(vals,colors)):
                p = row.cells[i].paragraphs[0]
                p.paragraph_format.left_indent  = Pt(4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                add_run(p, val, size=10, color=col)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    og_tags = D.get('ogTags',[])
    if og_tags:
        check_row(doc, 'p', 'Open Graph Tags — configured',
                  '  ·  '.join(f"{t.get('t')}: {t.get('v','')[:40]}" for t in og_tags[:3]))
    else:
        check_row(doc, 'w', 'Open Graph Tags — missing')
    doc.add_paragraph()

    # ── LOCAL SEO ─────────────────────────────────────────────────────────────
    section_heading(doc, '📍  Local SEO', '0E6655')
    check_row(doc, 'p' if loc.get('hasAddress') else 'w',
              f'Address & Phone — {"visible on page" if loc.get("hasAddress") else "not found"}',
              f'{loc.get("phone","")}  ·  {loc.get("addr","")}' if loc.get('hasAddress') else '')
    check_row(doc, 'p' if loc.get('localSchema') else 'w',
              f'Local Business Schema — {"found" if loc.get("localSchema") else "not found"}')
    check_row(doc, 'p' if gbp.get('found') else 'w',
              f'Google Business Profile — {"found" if gbp.get("found") else "not found"}',
              f'{gbp.get("name","")}  ·  {gbp.get("addr","")}' if gbp.get('found') else '')
    rating = rev.get('rating',0)
    if rating:
        stars = '★'*round(rating)+'☆'*(5-round(rating))
        check_row(doc, 'p' if rating>=4 else 'w',
                  f'Google Reviews — {rating} {stars} ({rev.get("count",0)} reviews)')
    doc.add_paragraph()

    # ── TECHNOLOGY ────────────────────────────────────────────────────────────
    section_heading(doc, '⚙️  Technology Results', '2C3E50')
    tech_list = tech.get('list',[])
    if tech_list:
        tech_tbl = doc.add_table(rows=1, cols=2)
        tech_tbl.style = 'Table Grid'
        set_table_borders(tech_tbl, 'DDDDDD')
        for i,hd in enumerate(['Technology','Version']):
            cell = tech_tbl.rows[0].cells[i]
            set_cell_bg(cell,'F8F9FB')
            hp = cell.paragraphs[0]
            hp.paragraph_format.left_indent = Pt(4)
            hp.paragraph_format.space_before = Pt(3)
            hp.paragraph_format.space_after  = Pt(3)
            add_run(hp, hd, bold=True, size=8, color=RGBColor(0x95,0xA5,0xA6))
        for t in tech_list:
            row = tech_tbl.add_row()
            for i,val in enumerate([t.get('name',''), t.get('ver','')]):
                p = row.cells[i].paragraphs[0]
                p.paragraph_format.left_indent  = Pt(4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                add_run(p, val, size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    check_row(doc, 'p' if tech.get('dmarc') else 'f',
              f'DMARC Record — {"found" if tech.get("dmarc") else "not found"}',
              tech.get('dmarcDesc',''))
    check_row(doc, 'p' if tech.get('spf') else 'w',
              f'SPF Record — {"found" if tech.get("spf") else "not found"}',
              tech.get('spfRecord',''))
    doc.add_paragraph()

    # ── RECOMMENDATIONS ───────────────────────────────────────────────────────
    section_heading(doc, '🎯  Priority Recommendations', '935116')
    for i, r in enumerate(recs):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        add_run(p, f'{r.get("priority",i+1)}.  ', bold=True, size=12,
                color=RGBColor(0xB7,0x77,0x0D))
        add_run(p, r.get('title',''), bold=True, size=12)
        pd = doc.add_paragraph()
        pd.paragraph_format.left_indent  = Cm(1.2)
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after  = Pt(8)
        add_run(pd, r.get('detail',''), size=11,
                color=RGBColor(0x5D,0x40,0x37))
    doc.add_paragraph()

    # ── FOOTER ────────────────────────────────────────────────────────────────
    ft_tbl = doc.add_table(rows=1, cols=1)
    ft_tbl.style = 'Table Grid'
    set_table_borders(ft_tbl, 'EEEEEE')
    fc = ft_tbl.rows[0].cells[0]
    set_cell_bg(fc, 'F8F9FB')
    fp = fc.paragraphs[0]
    fp.paragraph_format.space_before = Pt(8)
    fp.paragraph_format.space_after  = Pt(8)
    fp.paragraph_format.left_indent  = Pt(6)
    add_run(fp, f'SEO Audit for {domain}  ·  Powered by Claude AI  ·  {date}',
            size=9, color=RGBColor(0x95,0xA5,0xA6))

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Send email ────────────────────────────────────────────────────────────────

def send_email_with_attachment(to_addr, subject, html_body, text_body,
                                docx_bytes, filename):
    host     = os.environ.get("SMTP_HOST", "")
    port     = int(os.environ.get("SMTP_PORT", "587"))
    user     = os.environ.get("SMTP_USER", "")
    pwd      = os.environ.get("SMTP_PASS", "")
    from_hdr = os.environ.get("SMTP_FROM", user)

    if not host or not user or not pwd:
        raise ValueError("SMTP not configured. Set SMTP_HOST, SMTP_USER, SMTP_PASS in Vercel env vars.")

    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"]    = from_hdr
    msg["To"]      = to_addr

    alt = MIMEMultipart("alternative")
    alt.attach(MIMEText(text_body, "plain"))
    alt.attach(MIMEText(html_body, "html"))
    msg.attach(alt)

    # DOCX attachment
    att = MIMEBase("application",
                   "vnd.openxmlformats-officedocument.wordprocessingml.document")
    att.set_payload(docx_bytes)
    encoders.encode_base64(att)
    att.add_header("Content-Disposition", "attachment", filename=filename)
    msg.attach(att)

    ctx = ssl.create_default_context()
    with smtplib.SMTP(host, port) as server:
        server.ehlo()
        server.starttls(context=ctx)
        server.login(user, pwd)
        server.sendmail(user, to_addr, msg.as_string())


# ── Handler ───────────────────────────────────────────────────────────────────

class handler(BaseHTTPRequestHandler):

    def do_POST(self):
        try:
            length   = int(self.headers.get("Content-Length", 0))
            body     = json.loads(self.rfile.read(length))
            job_id   = body.get("job_id", "").strip()
            to_email = body.get("email",  "").strip()

            if not to_email or not valid_email(to_email):
                return self._json(400, {"error": "Please enter a valid email address."})

            record = store_get(job_id)
            if not record or record.get("status") != "done":
                return self._json(404, {"error": "Report not found. Please run the audit again."})

            D      = record.get("data", {})
            name   = record.get("name", "there")
            domain = D.get("domain", "your-website")

            from datetime import datetime
            date     = datetime.now().strftime("%d %B %Y")
            filename = f"seo-report-{domain}-{datetime.now().strftime('%Y%m%d')}.docx"

            html_body  = build_summary_html(D, name, to_email)
            text_body  = build_summary_text(D, name)
            docx_bytes = build_docx(D, name, to_email, date)

            send_email_with_attachment(
                to_email,
                f"Your SEO Audit Report for {domain}",
                html_body,
                text_body,
                docx_bytes,
                filename
            )

            self._json(200, {"ok": True, "message": f"Report sent to {to_email}"})

        except ValueError as e:
            self._json(400, {"error": str(e)})
        except Exception as e:
            self._json(500, {"error": f"Failed to send: {str(e)}"})

    def do_OPTIONS(self):
        self.send_response(204)
        self._cors()
        self.end_headers()

    def _json(self, code, data):
        body = json.dumps(data).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self._cors()
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin",  "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def log_message(self, *a): pass
